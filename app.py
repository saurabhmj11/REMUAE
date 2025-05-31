import streamlit as st
import subprocess
import os
import tempfile
import numpy as np
import soundfile as sf
import shutil
import pathlib
import io # For handling in-memory files like BytesIO
import datetime # <--- FIX: Added/Moved this import here

# Hugging Face and Faster Whisper imports
from faster_whisper import WhisperModel
from transformers import pipeline

# New imports for document generation
from docx import Document
from docx.shared import Inches
from fpdf import FPDF # fpdf2 is imported as FPDF

# --- Configuration & Constants ---
ACCENT_LABEL_MAP = {
    'en-US': 'American English (US)',
    'en-AU': 'Australian English',
    'en-CA': 'Canadian English',
    'en-IN': 'Indian English',
    'en-GB': 'British English (GB)',
    'en-NZ': 'New Zealand English',
    # Add more if the model supports them or you interpret them
}

# --- Model Loading (Cached) ---
@st.cache_resource
def load_whisper_model():
    """Load the faster-whisper model once."""
    st.info("Loading Faster Whisper model (this happens once)...")
    return WhisperModel("base.en", device="cpu", compute_type="int8")

@st.cache_resource
def load_accent_model():
    """Load the accent classification model once."""
    st.info("Loading accent classification model (this happens once)...")
    return pipeline(
        "audio-classification",
        model="lhotse-speech/wav2vec2-base-ft-cv13-accent-classifier",
        trust_remote_code=True
    )

# --- Audio Extraction Functions ---
def extract_audio_from_url(video_url, output_dir):
    """
    Extracts audio from a given video URL using yt-dlp.
    Returns the path to the extracted audio file.
    """
    audio_path = os.path.join(output_dir, "extracted_audio.wav")
    status_text = st.empty()

    try:
        status_text.info(f"⏳ Extracting audio from URL: {video_url}...")
        command = [
            "yt-dlp",
            "-x",
            "--audio-format", "wav",
            "-o", str(pathlib.Path(audio_path).with_suffix('')),
            video_url
        ]
        subprocess.run(
            command,
            check=True,
            capture_output=True,
            text=True
        )
        status_text.success("✅ Audio extracted successfully!")
        return audio_path
    except subprocess.CalledProcessError as e:
        status_text.error(f"❌ Error extracting audio: {e.stderr}. Please check the URL or try again later.")
        st.code(e.stderr)
        return None
    except FileNotFoundError:
        status_text.error("❌ `yt-dlp` not found. Please ensure it's installed and in your system's PATH.")
        return None
    except Exception as e:
        status_text.error(f"❌ An unexpected error occurred during audio extraction: {e}")
        return None

def process_uploaded_audio(uploaded_file, output_dir):
    """
    Saves the uploaded audio file to a temporary location.
    """
    file_extension = uploaded_file.name.split('.')[-1].lower()
    temp_audio_path = os.path.join(output_dir, f"uploaded_audio.{file_extension}")
    try:
        with open(temp_audio_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        st.success(f"✅ Audio file '{uploaded_file.name}' uploaded and saved successfully.")
        return temp_audio_path
    except Exception as e:
        st.error(f"❌ Error saving uploaded audio: {e}")
        return None

# --- Accent Analysis & Transcription ---
def analyze_accent(audio_file_path):
    """
    Analyzes the accent using the pre-trained Hugging Face model.
    """
    try:
        accent_classifier = load_accent_model()
        results = accent_classifier(audio_file_path)

        if results:
            best_match = results[0]
            accent_label = ACCENT_LABEL_MAP.get(best_match['label'], best_match['label'].replace('en-', '').capitalize() + " English")
            confidence = best_match['score'] * 100

            english_speaking_candidate_confidence = confidence

            return accent_label, confidence, english_speaking_candidate_confidence, results
        else:
            return "Unable to classify", 0, 0, []
    except Exception as e:
        # This error is specifically for the Hugging Face model loading
        st.error(f"❌ Error during accent analysis: {e}")
        return "Error", 0, 0, []

def transcribe_audio(audio_file_path):
    """
    Transcribes audio using Faster Whisper.
    """
    try:
        st.info("⏳ Transcribing audio (this may take a moment for longer videos)...")
        model = load_whisper_model()
        segments, info = model.transcribe(audio_file_path, beam_size=5, language="en")

        full_transcription = ""
        transcription_lines = []
        for i, segment in enumerate(segments):
            full_transcription += segment.text + " "
            transcription_lines.append(f"**[{int(segment.start)}s - {int(segment.end)}s]**: {segment.text.strip()}")
            if i < 50: # Limit displaying too many lines in UI
                pass # Already handled by Streamlit's expander

        if not full_transcription.strip():
            st.warning("🧐 No clear English speech was transcribed. This might indicate silence, background noise, or a non-English language.")
            return "", []

        st.success(f"✅ Transcription complete. Detected language: {info.language} (Confidence: {info.language_probability:.2%})")
        return full_transcription.strip(), transcription_lines
    except Exception as e:
        st.error(f"❌ Error during transcription: {e}")
        return "", []

# --- Document Generation Functions ---
def generate_report_markdown(accent_data, transcription_data):
    """
    Generates a Markdown string containing the full report.
    """
    report_md = f"""
# REM Waste AI Accent Analysis Report

## 📊 Analysis Results

*   **Detected English Accent:** {accent_data['accent']}
*   **Accent Classification Confidence:** {accent_data['confidence_accent']:.1f}%
*   **Overall English Candidate Assessment:** {accent_data['confidence_english_speaking']:.1f}%
    (Based on successful English transcription and accent classification confidence)

---

## 📈 Raw Accent Scores
"""
    if accent_data['raw_results']:
        for res in accent_data['raw_results']:
            label_display = ACCENT_LABEL_MAP.get(res['label'], res['label'])
            report_md += f"*   **{label_display}:** {res['score']:.2%}\n"
    else:
        report_md += "*   No detailed accent scores available.\n"

    report_md += """
---

## 📝 Summary & Explanation

The tool first extracted audio from the provided input (URL or uploaded file).
The audio was then processed using an optimized **Faster Whisper** model to confirm English speech and generate a transcription.

Finally, a specialized AI model trained on various English accents
(including **American, British, Australian, Canadian, Indian, and New Zealand English**)
analyzed the acoustic features of the speech to classify the accent.

The `Accent Classification Confidence` indicates how certain the model is about its primary prediction.
The `Overall English Candidate Assessment` is a general indicator of how confidently the system could process and classify English speech from the input.

---

## 🎙️ Full Transcription

"""
    if transcription_data['formatted_lines']:
        for line in transcription_data['formatted_lines']:
            # Replace markdown bolding for plain text/docx conversion convenience
            plain_line = line.replace('**', '').replace(':', ':', 1) # Keep first colon
            report_md += f"- {plain_line}\n"
    else:
        report_md += "No transcription available.\n"

    report_md += f"\n---"
    report_md += f"\nGenerated on: {st.session_state.analysis_timestamp.strftime('%Y-%m-%d %H:%M:%S')} (UTC)"


    return report_md

def generate_docx(report_md_content):
    """Generates a .docx file from the report content."""
    document = Document()
    document.add_heading('REM Waste AI Accent Analysis Report', level=0)

    # Simple parsing of Markdown. For a truly rich Markdown document,
    # you'd need a more sophisticated Markdown parser.
    # This just splits by lines and handles basic headers/lists.
    lines = report_md_content.split('\n')
    for line in lines:
        if line.startswith('# '):
            level = line.count('#')
            document.add_heading(line.lstrip('# ').strip(), level=min(level, 4)) # Cap at level 4 for docx
        elif line.startswith('*   ') or line.startswith('- '):
            document.add_paragraph(line.lstrip('* - ').strip(), style='List Bullet')
        elif line.strip() == '---':
            document.add_paragraph().add_run().add_break() # Add a line break for horizontal rule
        elif line.strip(): # Avoid empty paragraphs
            document.add_paragraph(line.strip())

    bio = io.BytesIO()
    document.save(bio)
    return bio.getvalue()

def generate_pdf(report_md_content):
    """Generates a basic .pdf file from the report content using fpdf2."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Basic parsing for PDF (fpdf2 does not natively render Markdown)
    for line in report_md_content.split('\n'):
        if line.startswith('# '): # Headers
            level = line.count('#')
            pdf.set_font("Arial", "B", size=20 - level * 2) # Reduce font size for lower levels
            pdf.ln(h=pdf.font_size*1.5) # Add extra space before header
            pdf.multi_cell(0, 10, line.lstrip('# ').strip(), align='C')
            pdf.ln(h=pdf.font_size*0.5)
            pdf.set_font("Arial", size=12) # Reset font
        elif line.startswith('*   ') or line.startswith('- '): # List items (simple bullet)
            pdf.multi_cell(0, 10, f"• {line.lstrip('* - ').strip()}")
        elif line.strip() == '---': # Horizontal rule
            pdf.ln(h=5)
            pdf.line(pdf.get_x(), pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(h=5)
        elif line.strip(): # Regular paragraph
            pdf.multi_cell(0, 10, line.strip())

    bio = io.BytesIO()
    pdf.output(bio)
    return bio.getvalue()


# --- Streamlit UI ---
st.set_page_config(page_title="REM Waste Accent Analyzer", layout="centered")

st.title("🗣️ REM Waste AI Accent Analyzer")
st.markdown(
    """
    Welcome to the REM Waste AI Accent Analyzer! This tool helps evaluate spoken English for hiring purposes
    by analyzing the speaker's accent from a provided video URL or audio file.

    It extracts audio, transcribes the speech, and then classifies the English accent.
    """
)

# Input method selection
input_method = st.radio(
    "Choose input method:",
    ("Public Video URL", "Upload Audio File (.wav, .mp3)"),
    index=0, # Default to URL
    key="input_method"
)

video_url = None
uploaded_file = None

if input_method == "Public Video URL":
    video_url = st.text_input(
        "Enter Public Video URL (e.g., Loom, direct MP4, YouTube):",
        key="video_url_input",
        placeholder="e.g., https://www.youtube.com/watch?v=dQw4w9WgXcQ"
    )
    if video_url and not (video_url.startswith("http://") or video_url.startswith("https://")):
        st.warning("Please enter a valid URL starting with http:// or https://")

elif input_method == "Upload Audio File (.wav, .mp3)":
    uploaded_file = st.file_uploader(
        "Upload an Audio File (.wav, .mp3):",
        type=["wav", "mp3"],
        help="Supported audio formats are WAV and MP3."
    )

st.markdown("---")

# Main Analysis Button
if st.button("🚀 Analyze Accent", type="primary", use_container_width=True):
    if not video_url and not uploaded_file:
        st.warning("Please provide a video URL or upload an audio file to proceed.")
    else:
        temp_dir = tempfile.mkdtemp()
        audio_file_path = None
        results_area = st.empty()


        try:
            with results_area.container(): # All processing output goes in this container
                st.session_state.analysis_timestamp = st.session_state.get('analysis_timestamp', None)
                if video_url:
                    audio_file_path = extract_audio_from_url(video_url, temp_dir)
                elif uploaded_file:
                    audio_file_path = process_uploaded_audio(uploaded_file, temp_dir)

                if audio_file_path and os.path.exists(audio_file_path):
                    # 1. Transcribe audio to confirm English speaking
                    transcription_text, formatted_transcription_lines = transcribe_audio(audio_file_path)

                    if not transcription_text or "No clear English speech" in st.session_state.get('_log_messages', ''): # Check for warning message too
                        st.error("🚫 Cannot proceed with accent analysis as no significant English speech was detected.")
                    else:
                        # 2. Analyze accent
                        st.subheader("📊 Accent Analysis Results")
                        accent, confidence_accent, confidence_english_speaking, raw_results = analyze_accent(audio_file_path)

                        col1, col2 = st.columns(2)
                        with col1:
                            st.metric(label="Detected English Accent", value=accent)
                        with col2:
                            st.metric(label="Accent Classification Confidence", value=f"{confidence_accent:.1f}%")

                        st.metric(label="Overall English Candidate Assessment", value=f"{confidence_english_speaking:.1f}%",
                                  help="This score reflects the confidence that the candidate is speaking English based on successful transcription and the accent classifier's certainty.")

                        st.success("✨ Analysis complete!")
                        st.balloons() # Little celebration

                        st.subheader("📈 Raw Accent Scores")
                        st.markdown("_Shows the probability for each supported English accent_")
                        if raw_results:
                            for res in raw_results:
                                label_display = ACCENT_LABEL_MAP.get(res['label'], res['label'])
                                st.progress(res['score'], text=f"**{label_display}:** {res['score']:.2%}")
                        else:
                            st.info("No detailed accent scores available.")

                        st.subheader("📝 Summary & Explanation")
                        st.markdown(f"""
                        The tool first extracted audio from the provided input (URL or uploaded file).
                        The audio was then processed using an optimized **Faster Whisper** model to confirm English speech and generate a transcription.

                        Finally, a specialized AI model trained on various English accents
                        (including **American, British, Australian, Canadian, Indian, and New Zealand English**)
                        analyzed the acoustic features of the speech to classify the accent.

                        The `Accent Classification Confidence` indicates how certain the model is about its primary prediction.
                        The `Overall English Candidate Assessment` is a general indicator of how confidently the system could process and classify English speech from the input.
                        """)


                        st.subheader("🎙️ Full Transcription")
                        with st.expander("Click to view full transcription", expanded=False):
                            if formatted_transcription_lines:
                                for line in formatted_transcription_lines:
                                    st.markdown(line)
                            else:
                                st.write("No transcription available.")

                        # Store results in session state for download buttons to access
                        st.session_state.accent_data = {
                            'accent': accent,
                            'confidence_accent': confidence_accent,
                            'confidence_english_speaking': confidence_english_speaking,
                            'raw_results': raw_results
                        }
                        st.session_state.transcription_data = {
                            'full_text': transcription_text,
                            'formatted_lines': formatted_transcription_lines
                        }
                        st.session_state.analysis_timestamp = datetime.datetime.utcnow() # Store timestamp

                elif not audio_file_path:
                    st.error("Audio extraction/processing failed. Please check your input or try again.")

        finally:
            # Clean up temporary files regardless of success or failure
            if temp_dir and os.path.exists(temp_dir):
                st.info("🗑️ Cleaning up temporary files...")
                try:
                    shutil.rmtree(temp_dir)
                    st.success("✅ Temporary files cleaned up.")
                except Exception as e:
                    st.warning(f"⚠️ Could not fully clean up temporary directory {temp_dir}: {e}")


# --- Download Options (only appear after successful analysis) ---
if st.session_state.get('accent_data') and st.session_state.get('transcription_data'):
    st.markdown("---")
    st.subheader("⬇️ Download Analysis Report")

    report_md_content = generate_report_markdown(st.session_state.accent_data, st.session_state.transcription_data)

    col_dl1, col_dl2, col_dl3 = st.columns(3)

    with col_dl1:
        st.download_button(
            label="📄 Download as Text (.txt)",
            data=report_md_content.encode("utf-8"),
            file_name=f"rem_waste_accent_report_{st.session_state.analysis_timestamp.strftime('%Y%m%d_%H%M%S')}.txt",
            mime="text/plain",
            help="Download a plain text version of the full report."
        )

    with col_dl2:
        st.download_button(
            label="📝 Download as Word (.docx)",
            data=generate_docx(report_md_content),
            file_name=f"rem_waste_accent_report_{st.session_state.analysis_timestamp.strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            help="Download a Microsoft Word document version of the full report."
        )

    with col_dl3:
        st.download_button(
            label="📥 Download as PDF (.pdf)",
            data=generate_pdf(report_md_content),
            file_name=f"rem_waste_accent_report_{st.session_state.analysis_timestamp.strftime('%Y%m%d_%H%M%S')}.pdf",
            mime="application/pdf",
            help="Download a basic PDF version of the full report. (Note: Formatting may be basic)."
        )

st.markdown("---")
st.markdown("💡 **Tip**: For best results, ensure the audio has clear, singular speech without significant background noise.")
st.caption("Developed for REM Waste Hiring Decisions as a practical challenge solution.")
st.markdown("Please note: Accent classification models may sometimes misinterpret accents, especially with short or ambiguous speech. This tool is for indicative purposes.")